import uuid
import os
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.config import settings

class DocxGenerator:
    def __init__(self):
        self.storage_path = settings.storage_path
        os.makedirs(self.storage_path, exist_ok=True)
    
    def generate(self, resume_data: dict) -> str:
        doc = Document()
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Name
        name = doc.add_paragraph()
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name.add_run(resume_data.get("name", ""))
        run.bold = True
        run.font.size = Pt(16)
        
        # Contact
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run(f"{resume_data.get('email', '')} | {resume_data.get('phone', '')} | {resume_data.get('location', '')}")
        contact.runs[0].font.size = Pt(10)
        
        # Summary
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(resume_data.get("summary", ""))
        
        # Skills
        doc.add_heading('Skills', level=1)
        doc.add_paragraph(", ".join(resume_data.get("skills", [])))
        
        # Experience
        doc.add_heading('Experience', level=1)
        for exp in resume_data.get("experience", []):
            p = doc.add_paragraph()
            run = p.add_run(f"{exp.get('title', '')} | {exp.get('company', '')}")
            run.bold = True
            
            dates = doc.add_paragraph(exp.get("dates", ""))
            dates.runs[0].italic = True
            
            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style='List Bullet')
        
        # Education
        if resume_data.get("education"):
            doc.add_heading('Education', level=1)
            for edu in resume_data.get("education", []):
                doc.add_paragraph(f"{edu.get('degree', '')} - {edu.get('school', '')}")
        
        # Save
        file_id = str(uuid.uuid4())
        filename = f"{file_id}.docx"
        filepath = os.path.join(self.storage_path, filename)
        doc.save(filepath)
        
        return filepath