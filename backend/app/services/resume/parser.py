import re
import json
from typing import Optional
from io import BytesIO
import fitz  # PyMuPDF
from docx import Document


class ResumeParser:
    """Parse uploaded resume files (PDF/DOCX) into structured text."""
    
    def parse(self, file_bytes: bytes, filename: str) -> str:
        """Parse a resume file and return raw text."""
        if filename.lower().endswith('.pdf'):
            return self._parse_pdf(file_bytes)
        elif filename.lower().endswith('.docx'):
            return self._parse_docx(file_bytes)
        elif filename.lower().endswith('.doc'):
            return self._parse_docx(file_bytes)  # python-docx handles .doc too
        elif filename.lower().endswith('.txt'):
            return file_bytes.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported file format: {filename}")
    
    def _parse_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF."""
        text_parts = []
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        for page in doc:
            text = page.get_text("text")
            if text.strip():
                text_parts.append(text)
        
        doc.close()
        return "\n".join(text_parts)
    
    def _parse_docx(self, file_bytes: bytes) -> str:
        """Extract text from DOCX."""
        doc = Document(BytesIO(file_bytes))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n".join(text_parts)